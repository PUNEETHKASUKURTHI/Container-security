import subprocess
import json
from kubernetes import client, config
from prometheus_client import start_http_server
from docx import Document
from docx.shared import Pt


# Load Kubernetes configuration
config.load_kube_config()

# Create Kubernetes API client
api_client = client.CoreV1Api()


def scan_container_images(namespace):
    # Get all pod names in the specified namespace
    pods = api_client.list_namespaced_pod(namespace)

    # Create a new Word document
    doc = Document()
    doc.add_heading('Container Security Scan Results', level=1)

    for pod in pods.items:
        # Get container images from pod
        container_images = [container.image for container in pod.spec.containers]

        for image in container_images:
            try:
                # Run Trivy scan command
                scan_command = ['trivy', 'image', '--format', 'json', image]
                scan_result = subprocess.run(scan_command, capture_output=True, text=True, check=True)

                # Print the raw Trivy scan output
                print(scan_result.stdout)

                # Process Trivy scan output
                vulnerabilities = json.loads(scan_result.stdout)
                if vulnerabilities:
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'

                    # Add table headers
                    headers = table.rows[0].cells
                    headers[0].text = 'Image'
                    headers[1].text = 'Vulnerability ID'
                    headers[2].text = 'Severity'
                    headers[3].text = 'Description'

                    for vulnerability in vulnerabilities:
                        if isinstance(vulnerability, dict):
                            # Add a new row to the table
                            row = table.add_row().cells
                            row[0].text = image
                            row[1].text = vulnerability.get('VulnerabilityID', '')
                            row[2].text = vulnerability.get('Severity', '')
                            row[3].text = vulnerability.get('Description', '')

            except subprocess.CalledProcessError as e:
                print(f"Error scanning image {image}: {e.stderr}")

    # Save the Word document
    doc.save('scan_results.docx')


if __name__ == '__main__':
    # Example usage: Scan container images in "default" namespace
    start_http_server(8000)
    scan_container_images("default")

